#!/usr/bin/env python3
"""
UCALP - Formateador de Documentos Académicos
Departamento de Educación a Distancia - v1.1

Uso:
    python ucalp_converter.py archivo.docx [opciones]
    python ucalp_converter.py carpeta/ [opciones]

Opciones:
    --facultad "Facultad de Humanidades"
    --carrera "Licenciatura en Filosofía"
    --asignatura "Filosofía social"
    --unidad "Unidad 1"
    --logo ruta/al/logo.png
    --salida carpeta_salida/
    --formatos pdf,html,docx

Requisitos:
    pip install python-docx reportlab mammoth Pillow
"""

import os, sys, re, argparse, base64
from pathlib import Path
from datetime import datetime

# ===== COLORES INSTITUCIONALES POR FACULTAD =====
# Cada facultad tiene su propio color distintivo.
# La paleta completa se deriva del color primario de la facultad.
FACULTY_COLORS = {
    'Facultad de Humanidades': {
        'primary': (11, 94, 126),     # Azul UCALP
        'dark':    (8, 61, 86),
        'light':   (228, 240, 247),
        'accent':  (26, 127, 170),
    },
    'Facultad de Ciencias Exactas e Ingeniería': {
        'primary': (180, 40, 40),     # Rojo
        'dark':    (120, 20, 20),
        'light':   (250, 225, 225),
        'accent':  (210, 70, 70),
    },
    'Facultad de Odontología': {
        'primary': (180, 140, 30),    # Mostaza
        'dark':    (120, 90, 10),
        'light':   (250, 240, 205),
        'accent':  (210, 170, 60),
    },
    'Facultad de Ciencias Económicas y Sociales': {
        'primary': (100, 55, 145),    # Violeta
        'dark':    (65, 30, 100),
        'light':   (235, 220, 255),
        'accent':  (130, 80, 180),
    },
    'Facultad de Derecho y Ciencias Políticas': {
        'primary': (195, 100, 20),    # Naranja
        'dark':    (140, 65, 10),
        'light':   (255, 235, 205),
        'accent':  (225, 130, 50),
    },
    'Facultad de Arquitectura y Diseño': {
        'primary': (35, 140, 75),     # Verde
        'dark':    (20, 90, 45),
        'light':   (210, 245, 225),
        'accent':  (55, 175, 100),
    },
    'Facultad de Ciencias de la Salud': {
        'primary': (185, 30, 110),    # Magenta/Fucsia
        'dark':    (120, 15, 70),
        'light':   (255, 215, 235),
        'accent':  (215, 60, 140),
    },
}

# Colores comunes para todos (no dependen de la facultad)
UCALP_COMMON = {
    'gold':  (200, 164, 78),
    'warm':  (245, 240, 232),
    'text':  (44, 62, 80),
    'muted': (107, 123, 141),
    'white': (255, 255, 255),
}

# Fallback al azul UCALP si la facultad no se encuentra
DEFAULT_FACULTY = 'Facultad de Humanidades'

def get_faculty_palette(facultad):
    """Retorna la paleta completa (primary+dark+light+accent + comunes) para la facultad dada."""
    base = FACULTY_COLORS.get(facultad, FACULTY_COLORS[DEFAULT_FACULTY])
    return {**base, **UCALP_COMMON}

def hex_color(name, palette=None):
    if palette is None:
        palette = get_faculty_palette(DEFAULT_FACULTY)
    r, g, b = palette[name]
    return f'#{r:02x}{g:02x}{b:02x}'

def get_logo_base64(logo_path):
    if logo_path and os.path.exists(logo_path):
        with open(logo_path, 'rb') as f:
            return base64.b64encode(f.read()).decode()
    return None

def escape_html(text):
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')


# ===== EXTRACCION =====
def extract_content_from_docx(filepath):
    import docx
    doc = docx.Document(filepath)
    content = {'paragraphs': [], 'headings': [], 'images': [], 'tables': []}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            content['paragraphs'].append({'type': 'empty', 'text': ''})
            continue

        style_name = para.style.name.lower() if para.style else ''

        if 'heading' in style_name or 'título' in style_name:
            level = 1
            if '2' in style_name: level = 2
            elif '3' in style_name: level = 3
            content['paragraphs'].append({'type': 'heading', 'level': level, 'text': text})
            content['headings'].append({'level': level, 'text': text})
        elif re.match(r'^[a-f]\)\s+', text):
            content['paragraphs'].append({'type': 'heading', 'level': 2, 'text': text})
            content['headings'].append({'level': 2, 'text': text})
        elif all(run.bold for run in para.runs if run.text.strip()) and len(text) < 200 and para.runs:
            content['paragraphs'].append({'type': 'heading', 'level': 2, 'text': text})
            content['headings'].append({'level': 2, 'text': text})
        elif text.startswith(('•', '-', '–', '▪')):
            content['paragraphs'].append({'type': 'list_item', 'text': text.lstrip('•-–▪ ')})
        else:
            runs = [{'text': r.text, 'bold': r.bold, 'italic': r.italic, 'underline': r.underline} for r in para.runs]
            content['paragraphs'].append({'type': 'paragraph', 'text': text, 'runs': runs})

    for table in doc.tables:
        content['tables'].append([[cell.text.strip() for cell in row.cells] for row in table.rows])

    return content


# ===== PDF =====
def generate_pdf(content, config, output_path, logo_path=None):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image

    palette = get_faculty_palette(config.get('facultad', DEFAULT_FACULTY))

    primary = HexColor(hex_color('primary', palette))
    dark = HexColor(hex_color('dark', palette))
    light = HexColor(hex_color('light', palette))
    gold = HexColor(hex_color('gold', palette))
    text_c = HexColor(hex_color('text', palette))
    muted = HexColor(hex_color('muted', palette))
    warm = HexColor(hex_color('warm', palette))

    styles = getSampleStyleSheet()
    s_body = ParagraphStyle('Body', parent=styles['Normal'], fontName='Helvetica', fontSize=11,
                            leading=17, textColor=text_c, alignment=TA_JUSTIFY, spaceAfter=8, firstLineIndent=24)
    s_h1 = ParagraphStyle('H1', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=18,
                          leading=22, textColor=dark, spaceBefore=20, spaceAfter=12)
    s_h2 = ParagraphStyle('H2', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=14,
                          leading=18, textColor=primary, spaceBefore=16, spaceAfter=8)
    s_h3 = ParagraphStyle('H3', parent=styles['Heading3'], fontName='Helvetica-BoldOblique', fontSize=12,
                          leading=16, textColor=dark, spaceBefore=12, spaceAfter=6)
    s_toc = ParagraphStyle('TOC', parent=styles['Normal'], fontName='Helvetica', fontSize=10,
                           leading=16, textColor=text_c, leftIndent=16, spaceBefore=2, spaceAfter=2)
    s_toc_title = ParagraphStyle('TOCTitle', parent=styles['Normal'], fontName='Helvetica-Bold',
                                 fontSize=11, leading=14, textColor=primary, spaceBefore=8, spaceAfter=8)
    s_list = ParagraphStyle('List', parent=s_body, leftIndent=24, bulletIndent=12, firstLineIndent=0)

    doc = SimpleDocTemplate(output_path, pagesize=A4, leftMargin=25*mm, rightMargin=25*mm,
                            topMargin=35*mm, bottomMargin=25*mm,
                            title=f"{config['asignatura']} - {config['unidad']}",
                            author="UCALP - Depto. Educación a Distancia")

    story = []

    # Header with logo
    if logo_path and os.path.exists(logo_path):
        logo_img = Image(logo_path, width=18*mm, height=18*mm)
        header_data = [[logo_img,
                        Paragraph('<b>UNIVERSIDAD CATÓLICA DE LA PLATA</b>',
                                  ParagraphStyle('hdr', fontName='Helvetica-Bold', fontSize=13, textColor=dark))]]
        ht = Table(header_data, colWidths=[22*mm, doc.width - 22*mm])
        ht.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                                ('LEFTPADDING', (0,0), (0,0), 0), ('LEFTPADDING', (1,0), (1,0), 6)]))
        story.append(ht)
    else:
        story.append(Paragraph('<b>UNIVERSIDAD CATÓLICA DE LA PLATA</b>',
                               ParagraphStyle('hdr', fontName='Helvetica-Bold', fontSize=14, textColor=dark)))

    story.append(Paragraph(f'<i>{config["facultad"]}</i>',
                           ParagraphStyle('fac', fontName='Helvetica-Oblique', fontSize=10, textColor=muted, spaceBefore=4)))
    story.append(Paragraph('Departamento de Educación a Distancia',
                           ParagraphStyle('dep', fontName='Helvetica', fontSize=9, textColor=muted, spaceBefore=2)))

    # Blue line
    lt = Table([['']], colWidths=[doc.width])
    lt.setStyle(TableStyle([('LINEBELOW', (0,0), (-1,-1), 2.5, primary),
                            ('TOPPADDING', (0,0), (-1,-1), 8), ('BOTTOMPADDING', (0,0), (-1,-1), 0)]))
    story.append(lt)
    story.append(Spacer(1, 6))

    # Meta bar
    meta = f"<b>Carrera:</b> {config['carrera']}  |  <b>Asignatura:</b> {config['asignatura']}  |  <b>{config['unidad']}</b>"
    story.append(Paragraph(meta, ParagraphStyle('meta', fontName='Helvetica', fontSize=9, textColor=muted,
                                                 backColor=warm, borderPadding=8, spaceAfter=16)))
    story.append(Spacer(1, 6))

    # Unit + title
    story.append(Paragraph(f"<font color='white'><b>&nbsp; {config['unidad'].upper()} &nbsp;</b></font>",
                           ParagraphStyle('badge', fontName='Helvetica-Bold', fontSize=9, backColor=primary,
                                          borderPadding=4, spaceAfter=6)))
    if content['headings']:
        story.append(Paragraph(content['headings'][0]['text'], s_h1))

    # TOC
    if len(content['headings']) > 1:
        story.append(Spacer(1, 8))
        story.append(Paragraph("CONTENIDOS", s_toc_title))
        tl = Table([['']], colWidths=[doc.width])
        tl.setStyle(TableStyle([('LINEABOVE', (0,0), (-1,0), 1, light)]))
        story.append(tl)
        for h in content['headings'][:10]:
            prefix = "◆ " if h['level'] <= 2 else "   ◇ "
            story.append(Paragraph(f"{prefix}{h['text']}", s_toc))
        story.append(tl)
        story.append(Spacer(1, 12))

    # Content
    for para in content['paragraphs']:
        if para['type'] == 'empty':
            story.append(Spacer(1, 4))
        elif para['type'] == 'heading':
            lvl = para.get('level', 2)
            story.append(Paragraph(para['text'], [s_h1, s_h1, s_h2, s_h3][min(lvl, 3)]))
        elif para['type'] == 'list_item':
            story.append(Paragraph(f"• {para['text']}", s_list))
        elif para['type'] == 'paragraph':
            if 'runs' in para and para['runs']:
                parts = []
                for r in para['runs']:
                    t = r['text'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    if r.get('bold') and r.get('italic'): t = f"<b><i>{t}</i></b>"
                    elif r.get('bold'): t = f"<b>{t}</b>"
                    elif r.get('italic'): t = f"<i>{t}</i>"
                    parts.append(t)
                story.append(Paragraph(''.join(parts), s_body))
            else:
                story.append(Paragraph(para['text'].replace('&','&amp;').replace('<','&lt;').replace('>','&gt;'), s_body))

    def add_footer(cv, d):
        cv.saveState()
        cv.setStrokeColor(HexColor(hex_color('light', palette)))
        cv.setLineWidth(1.5)
        cv.line(25*mm, 20*mm, A4[0]-25*mm, 20*mm)
        cv.setFont('Helvetica', 8)
        cv.setFillColor(muted)
        cv.drawString(25*mm, 15*mm, f"{config['asignatura']} — {config['unidad']}")
        cv.drawRightString(A4[0]-25*mm, 15*mm, f"UCALP | {config['facultad']}  —  Pág. {d.page}")
        cv.restoreState()

    try:
        doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
        return True
    except Exception as e:
        print(f"❌ Error PDF: {e}")
        return False


# ===== HTML =====
def generate_html(content, config, output_path, logo_path=None):
    palette = get_faculty_palette(config.get('facultad', DEFAULT_FACULTY))
    p_primary = hex_color('primary', palette)
    p_dark    = hex_color('dark',    palette)
    p_light   = hex_color('light',   palette)
    p_accent  = hex_color('accent',  palette)
    p_gold    = hex_color('gold',    palette)
    p_text    = hex_color('text',    palette)
    p_muted   = hex_color('muted',   palette)
    logo_html = ''
    if logo_path and os.path.exists(logo_path):
        b64 = get_logo_base64(logo_path)
        if b64:
            ext = Path(logo_path).suffix.lower().replace('.', '')
            mime = f'image/{ext}' if ext != 'jpg' else 'image/jpeg'
            logo_html = f'<img src="data:{mime};base64,{b64}" alt="UCALP">'

    parts = []
    for para in content['paragraphs']:
        if para['type'] == 'empty': continue
        elif para['type'] == 'heading':
            tag = f'h{min(para.get("level",2)+1, 4)}'
            parts.append(f'    <{tag}>{escape_html(para["text"])}</{tag}>')
        elif para['type'] == 'list_item':
            parts.append(f'    <li>{escape_html(para["text"])}</li>')
        elif para['type'] == 'paragraph':
            if 'runs' in para and para['runs']:
                fmt = []
                for r in para['runs']:
                    t = escape_html(r['text'])
                    if r.get('bold') and r.get('italic'): t = f'<strong><em>{t}</em></strong>'
                    elif r.get('bold'): t = f'<strong>{t}</strong>'
                    elif r.get('italic'): t = f'<em>{t}</em>'
                    fmt.append(t)
                parts.append(f'    <p>{"".join(fmt)}</p>')
            else:
                parts.append(f'    <p>{escape_html(para["text"])}</p>')

    body = '\n'.join(parts)

    toc = ''
    if content['headings']:
        items = '\n'.join([f'        <li>{escape_html(h["text"])}</li>' for h in content['headings'][:10]])
        toc = f'\n    <div class="toc"><h3>Contenidos de esta unidad</h3><ul>\n{items}\n      </ul></div>'

    html = f'''<!DOCTYPE html>
<html lang="es">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{escape_html(config["asignatura"])} - {escape_html(config["unidad"])} | UCALP</title>
  <style>
    :root {{ --primary:{p_primary}; --dark:{p_dark}; --light:{p_light}; --accent:{p_accent}; --gold:{p_gold}; --text:{p_text}; --muted:{p_muted}; }}
    * {{ margin:0; padding:0; box-sizing:border-box; }}
    body {{ font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif; color:var(--text); line-height:1.7; max-width:860px; margin:0 auto; padding:20px; background:#f5f7f9; }}
    .header {{ background:linear-gradient(135deg,var(--dark) 0%,var(--primary) 70%,var(--accent) 100%); color:white; padding:24px 28px; border-radius:12px 12px 0 0; display:flex; align-items:center; gap:18px; }}
    .header img {{ width:56px; height:56px; border-radius:50%; border:2px solid rgba(255,255,255,0.3); flex-shrink:0; }}
    .header-text h1 {{ font-size:20px; font-weight:700; margin:0 0 2px; }}
    .header-text .university {{ font-size:11px; text-transform:uppercase; letter-spacing:1px; opacity:0.75; margin-bottom:4px; }}
    .header-text p {{ font-size:13px; opacity:0.85; margin:0; }}
    .meta-bar {{ background:#f5f0e8; padding:10px 28px; font-size:12px; color:var(--muted); border-bottom:2px solid var(--gold); display:flex; gap:20px; flex-wrap:wrap; }}
    .meta-bar strong {{ color:var(--text); }}
    .content {{ background:white; padding:32px 28px; border-radius:0 0 12px 12px; box-shadow:0 2px 12px rgba(0,0,0,0.08); }}
    .toc {{ background:var(--light); border:1px solid #cde0eb; border-left:4px solid var(--primary); border-radius:0 8px 8px 0; padding:16px 20px; margin-bottom:24px; }}
    .toc h3 {{ font-size:13px; color:var(--primary); margin-bottom:10px; text-transform:uppercase; letter-spacing:0.5px; }}
    .toc ul {{ list-style:none; padding:0; }}
    .toc li {{ padding:4px 0; font-size:13px; border-bottom:1px solid rgba(11,94,126,0.1); }}
    .toc li::before {{ content:"◆ "; color:var(--gold); font-size:10px; }}
    h2 {{ font-size:20px; color:var(--dark); margin-top:28px; margin-bottom:12px; padding-bottom:8px; border-bottom:2px solid var(--light); }}
    h3 {{ font-size:16px; color:var(--primary); margin-top:20px; margin-bottom:10px; }}
    h4 {{ font-size:14px; color:var(--dark); margin-top:16px; margin-bottom:8px; }}
    p {{ font-size:14px; margin-bottom:12px; text-align:justify; }}
    em {{ color:var(--primary); }}
    ul, ol {{ margin:12px 0 12px 24px; }}
    li {{ margin-bottom:6px; font-size:14px; }}
    .footer {{ text-align:center; padding:16px; font-size:12px; color:var(--muted); margin-top:12px; }}
    @media (max-width:600px) {{ body {{ padding:10px; }} .header {{ padding:16px; flex-direction:column; text-align:center; }} .content {{ padding:20px 16px; }} }}
  </style>
</head>
<body>
  <div class="header">
    {logo_html}
    <div class="header-text">
      <div class="university">Universidad Católica de La Plata</div>
      <h1>{escape_html(config["asignatura"])}</h1>
      <p>{escape_html(config["facultad"])} · {escape_html(config["carrera"])}</p>
    </div>
  </div>
  <div class="meta-bar">
    <span><strong>Carrera:</strong> {escape_html(config["carrera"])}</span>
    <span><strong>{escape_html(config["unidad"])}</strong></span>
    <span><strong>UCALP</strong> — Depto. Educación a Distancia</span>
  </div>
  <div class="content">{toc}
{body}
  </div>
  <div class="footer">
    {escape_html(config["asignatura"])} — {escape_html(config["unidad"])} ·
    Universidad Católica de La Plata — {escape_html(config["facultad"])} ·
    Generado el {datetime.now().strftime("%d/%m/%Y")}
  </div>
</body>
</html>'''

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)
    return True


# ===== DOCX =====
def generate_docx(content, config, output_path, logo_path=None):
    import docx
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    palette = get_faculty_palette(config.get('facultad', DEFAULT_FACULTY))
    pr, pg, pb = palette['primary']
    dr, dg, db = palette['dark']

    C_PRIMARY = RGBColor(pr, pg, pb)
    C_DARK    = RGBColor(dr, dg, db)
    C_MUTED   = RGBColor(107, 123, 141)
    C_TEXT    = RGBColor(44, 62, 80)

    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = C_TEXT
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # Logo
    if logo_path and os.path.exists(logo_path):
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.add_run().add_picture(logo_path, width=Cm(2.2))

    # University name
    hp = doc.add_paragraph()
    hp.paragraph_format.space_before = Pt(4)
    r = hp.add_run('UNIVERSIDAD CATÓLICA DE LA PLATA')
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = C_DARK

    fp = doc.add_paragraph()
    r = fp.add_run(config['facultad'])
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = C_MUTED

    dp = doc.add_paragraph()
    dp.paragraph_format.space_after = Pt(2)
    r = dp.add_run('Departamento de Educación a Distancia')
    r.font.size = Pt(9); r.font.color.rgb = C_MUTED

    # Blue separator
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)
    pPr = sp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), f'{pr:02X}{pg:02X}{pb:02X}')
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Meta
    mp = doc.add_paragraph()
    mp.paragraph_format.space_before = Pt(4)
    mp.paragraph_format.space_after = Pt(12)
    for label, value in [('Carrera: ', config['carrera']), ('  |  Asignatura: ', config['asignatura']), ('  |  ', config['unidad'])]:
        r = mp.add_run(label); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = C_MUTED
        r = mp.add_run(value); r.font.size = Pt(9); r.font.color.rgb = C_TEXT

    # Unit
    up = doc.add_paragraph()
    up.paragraph_format.space_before = Pt(12)
    r = up.add_run(config['unidad'].upper())
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_PRIMARY

    if content['headings']:
        tp = doc.add_heading(content['headings'][0]['text'], level=1)
        for r in tp.runs: r.font.color.rgb = C_DARK; r.font.size = Pt(18)

    # TOC
    if len(content['headings']) > 1:
        th = doc.add_paragraph()
        th.paragraph_format.space_before = Pt(16)
        r = th.add_run('CONTENIDOS'); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_PRIMARY
        for h in content['headings'][:10]:
            tp = doc.add_paragraph()
            tp.paragraph_format.left_indent = Pt(12)
            tp.paragraph_format.space_before = Pt(2)
            tp.paragraph_format.space_after = Pt(2)
            prefix = '◆ ' if h['level'] <= 2 else '   ◇ '
            r = tp.add_run(f"{prefix}{h['text']}"); r.font.size = Pt(10); r.font.color.rgb = C_TEXT

    # Content
    for para in content['paragraphs']:
        if para['type'] == 'empty':
            doc.add_paragraph('')
        elif para['type'] == 'heading':
            level = min(para.get('level', 2), 3)
            h = doc.add_heading(para['text'], level=level)
            color = C_DARK if level != 2 else C_PRIMARY
            for r in h.runs: r.font.color.rgb = color
        elif para['type'] == 'list_item':
            doc.add_paragraph(para['text'], style='List Bullet')
        elif para['type'] == 'paragraph':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if 'runs' in para and para['runs']:
                for rr in para['runs']:
                    r = p.add_run(rr['text'])
                    r.font.size = Pt(11); r.font.color.rgb = C_TEXT
                    if rr.get('bold'): r.bold = True
                    if rr.get('italic'): r.italic = True; r.font.color.rgb = C_PRIMARY
                    if rr.get('underline'): r.underline = True
            else:
                r = p.add_run(para['text']); r.font.size = Pt(11); r.font.color.rgb = C_TEXT

    # Footer
    ft = doc.add_paragraph()
    ft.paragraph_format.space_before = Pt(24)
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ft.add_run(f"{config['asignatura']} — {config['unidad']} | UCALP — {config['facultad']}")
    r.font.size = Pt(8); r.font.color.rgb = C_MUTED

    doc.save(output_path)
    return True


# ===== MAIN =====
def main():
    parser = argparse.ArgumentParser(description='🎓 UCALP - Formateador de Documentos Académicos',
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument('input', help='Archivo .docx o carpeta')
    parser.add_argument('--facultad', default='Facultad de Humanidades')
    parser.add_argument('--carrera', default='Licenciatura en Filosofía')
    parser.add_argument('--asignatura', default='')
    parser.add_argument('--unidad', default='')
    parser.add_argument('--logo', default=None, help='Ruta al logo PNG')
    parser.add_argument('--salida', default='./output')
    parser.add_argument('--formatos', default='pdf,html,docx')

    args = parser.parse_args()
    input_path = Path(args.input)

    if input_path.is_dir():
        files = list(input_path.glob('*.docx'))
    elif input_path.is_file() and input_path.suffix == '.docx':
        files = [input_path]
    else:
        print(f"❌ No encontrado: {input_path}"); sys.exit(1)

    if not files:
        print("❌ No hay archivos .docx"); sys.exit(1)

    output_dir = Path(args.salida)
    output_dir.mkdir(parents=True, exist_ok=True)
    formatos = [f.strip().lower() for f in args.formatos.split(',')]
    logo_path = args.logo if args.logo and os.path.exists(args.logo) else None

    print("\n╔══════════════════════════════════════════════════════════╗")
    print("║   🎓 UCALP - Formateador de Documentos Académicos       ║")
    print("║   Departamento de Educación a Distancia                  ║")
    print("╚══════════════════════════════════════════════════════════╝")
    if logo_path: print(f"   🖼️  Logo: {logo_path}")
    print()

    for filepath in files:
        filename = filepath.stem
        asignatura = args.asignatura or filename.replace('_', ' ').replace('-', ' ')
        unidad = args.unidad
        if not unidad:
            m = re.search(r'[Uu]nidad[\s_]*(\d+)', filename)
            unidad = f"Unidad {m.group(1)}" if m else "Unidad 1"

        config = {'facultad': args.facultad, 'carrera': args.carrera, 'asignatura': asignatura, 'unidad': unidad}

        print(f"📄 {filepath.name}")
        print(f"   {config['facultad']} | {config['carrera']} | {config['asignatura']} | {config['unidad']}")

        content = extract_content_from_docx(str(filepath))
        print(f"   ✅ {len(content['paragraphs'])} párrafos, {len(content['headings'])} títulos")

        if 'pdf' in formatos:
            p = output_dir / f"{filename}_UCALP.pdf"
            if generate_pdf(content, config, str(p), logo_path): print(f"   📕 {p}")
        if 'html' in formatos:
            p = output_dir / f"{filename}_UCALP.html"
            if generate_html(content, config, str(p), logo_path): print(f"   🌐 {p}")
        if 'docx' in formatos:
            p = output_dir / f"{filename}_UCALP.docx"
            if generate_docx(content, config, str(p), logo_path): print(f"   📝 {p}")
        print()

    print(f"✨ Archivos en: {output_dir.absolute()}\n")

if __name__ == '__main__':
    main()
